import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json
from datetime import date
import os
from io import BytesIO
import pandas as pd

st.set_page_config(page_title="Duty Chart Generator", layout="wide")

# ---------------- Dark Theme CSS ----------------
dark_theme_css = """
<style>
[data-testid="stAppViewContainer"] {
    background-color: #0e1117;
    color: white;
}
[data-testid="stSidebar"] {
    background-color: #1c1e26;
}
.stTextInput > div > div > input,
.stTextArea > div > textarea,
.stSelectbox > div > div,
.stMultiSelect > div > div {
    background-color: #1c1e26 !important;
    color: white !important;
    border: 1px solid #444 !important;
    border-radius: 6px;
}
h1, h2, h3, h4, h5, h6, label, p {
    color: white !important;
}
.stDataFrame {
    background-color: #0e1117 !important;
    color: white !important;
}
.stButton > button {
    background-color: #2e3b4e !important;
    color: white !important;
    border-radius: 8px;
    border: 1px solid #444 !important;
}
.stButton > button:hover {
    background-color: #3e4f65 !important;
    border: 1px solid #666 !important;
}
</style>
"""
st.markdown(dark_theme_css, unsafe_allow_html=True)

# ---------------- Load Teams ----------------
DEFAULT_TEAMS = {
    "p_o": ["Ms. Sonia", "Priyanka Singh", "Ms. Reshmiya"],
    "audiologist": ["Mr. Vikram", "Mr. Aditya"],
    "edp": ["Mr. Paritosh", "Mr. Yogesh Sharma"],
    "spectacles": ["Spectacles Team"],
    "technician": ["Technician 1", "Technician 2"],
    "I/C": ["Mr Pankaj Dwivedi"]
}

def load_teams(path="teams.json"):
    """Load teams from teams.json or use defaults."""
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                for k, v in DEFAULT_TEAMS.items():
                    if k not in data:
                        data[k] = v
                return data
        except Exception as e:
            st.error(f"Error reading {path}: {e}")
            return DEFAULT_TEAMS.copy()
    return DEFAULT_TEAMS.copy()

teams = load_teams()

# ---------------- Session State ----------------
if "duties" not in st.session_state:
    st.session_state.duties = []

# ---------------- Camp Info Inputs ----------------
st.markdown("### 🏕️ Camp Information")

# 👉 New placeholder example
chart_title = st.text_input(
    "📝 Duty Chart For:",
    placeholder="DUTY CHART FOR"
)

col1, col2 = st.columns(2)
with col1:
    venue = st.text_input("📍 Enter Venue", placeholder="Fill Venue")
    sap_id = st.text_input("🆔 SAP ID", placeholder="Fill SAP ID")
with col2:
    camp_id = st.text_input("🎪 CAMP ID", placeholder="Fill CAMP ID")
    nob = st.text_input("🔢 NOB", placeholder="Fill NOB")

value = st.text_input("💰 VALUE", placeholder="Fill VALUE")

st.divider()

# ---------------- Main UI ----------------
st.title("📄 Duty Chart Generator")
st.markdown("Create and download a formatted **Duty Chart** for your camp with just a few clicks.")

# ---------------- Helper for Others ----------------
def multiselect_with_others(label, options, key_prefix):
    """Multiselect that allows an 'Others' option with text input."""
    selected = st.multiselect(label, options + ["Others"], key=f"multi_{key_prefix}")
    custom_name = ""
    if "Others" in selected:
        custom_name = st.text_input(
            f"Enter Other {label} Name(s)",
            key=f"other_{key_prefix}",
            placeholder="Type additional name(s), comma separated if multiple"
        )
        if custom_name:
            custom_names = [x.strip() for x in custom_name.split(",") if x.strip()]
            selected = [x for x in selected if x != "Others"] + custom_names
    return selected

# ---------------- Add Duty Row ----------------
with st.expander("➕ Add a Duty Row", expanded=True):
    col1, col2 = st.columns([1, 2])
    with col1:
        camp_date = st.date_input("📅 Select Date", value=date.today())
    with col2:
        team_headed = st.multiselect(
            "👩‍⚕️ Team Headed By",
            teams["I/C"] + teams["p_o"] + ["Others"],
            key="team_headed"
        )
        custom_team_headed = ""
        if "Others" in team_headed:
            custom_team_headed = st.text_input(
                "Enter Other Team Head Name(s)",
                key="other_team_headed",
                placeholder="Type additional name(s), comma separated if multiple"
            )
            if custom_team_headed:
                custom_names = [x.strip() for x in custom_team_headed.split(",") if x.strip()]
                team_headed = [x for x in team_headed if x != "Others"] + custom_names

    st.markdown("#### Select Team Members")
    col1, col2, col3, col4, col5 = st.columns(5)
    with col1:
        selected_po = multiselect_with_others("P&O", teams["p_o"], "po")
    with col2:
        selected_audiologist = multiselect_with_others("Audiologist", teams["audiologist"], "audiologist")
    with col3:
        selected_edp = multiselect_with_others("EDP", teams["edp"], "edp")
    with col4:
        selected_spectacles = multiselect_with_others("Spectacles", teams["spectacles"], "spectacles")
    with col5:
        selected_technician = multiselect_with_others("Technician", teams["technician"], "technician")

    reporting_time = st.text_area(
        "⏰ Reporting Time(s)",
        "Team report at camp venue on " + date.today().strftime("%d.%m.%Y"),
        help="Enter multiple reporting times if needed (each on a new line)"
    )

    if st.button("➕ Add Duty Row", type="primary", use_container_width=True):
        st.session_state.duties.append({
            "date": camp_date,
            "team_headed": team_headed,
            "p_o": selected_po,
            "audiologist": selected_audiologist,
            "edp": selected_edp,
            "spectacles": selected_spectacles,
            "technician": selected_technician,
            "reporting_time": reporting_time
        })
        st.toast("✅ Row added successfully!")

st.divider()

# ---------------- Preview ----------------
st.subheader("📋 Current Duty Rows")

if st.session_state.duties:
    preview_data = []
    for i, duty in enumerate(st.session_state.duties, start=1):
        preview_data.append({
            "S.No.": i,
            "Date": duty["date"].strftime('%d-%m-%Y'),
            "Headed By": ", ".join(duty["team_headed"]),
            "Team": ", ".join(
                duty["p_o"] + duty["audiologist"] + duty["edp"] +
                duty["spectacles"] + duty["technician"]
            ),
            "Reporting Time": duty["reporting_time"]
        })
    df = pd.DataFrame(preview_data)
    st.dataframe(df, use_container_width=True)

    for i in range(len(st.session_state.duties)):
        if st.button(f"❌ Remove Row {i+1}", key=f"remove_{i}"):
            st.session_state.duties.pop(i)
            st.rerun()
else:
    st.info("No rows added yet. Use the form above to add duty rows.")

st.divider()

# ---------------- Generate Word ----------------
def build_doc(duties, chart_title, venue, sap_id, camp_id, nob, value):
    doc = Document()

    # Heading
    if chart_title.strip():
        heading = doc.add_paragraph()
        hrun = heading.add_run(chart_title.upper())
        hrun.bold = True
        hrun.font.size = Pt(14)
        hrun.font.underline = True
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Sub details table (only show filled fields)
    sub_fields = []
    if sap_id.strip(): sub_fields.append(("SAP ID: " + sap_id, ""))
    if camp_id.strip():
        if len(sub_fields) == 0: sub_fields.append(("", "CAMP ID: " + camp_id))
        else: sub_fields[0] = (sub_fields[0][0], "CAMP ID: " + camp_id)
    if nob.strip(): sub_fields.append(("NOB = " + nob, ""))
    if value.strip():
        if len(sub_fields) == 3:
            sub_fields[2] = (sub_fields[2][0], "VALUE: " + value)
        elif len(sub_fields) == 2:
            sub_fields.append(("", "VALUE: " + value))
        elif len(sub_fields) == 1:
            sub_fields.append(("", "VALUE: " + value))

    if sub_fields:
        sub_table = doc.add_table(rows=len(sub_fields), cols=2)
        sub_table.style = None
        sub_table.autofit = True
        for r, (c1, c2) in enumerate(sub_fields):
            sub_table.rows[r].cells[0].text = c1
            sub_table.rows[r].cells[1].text = c2
            for cell in sub_table.rows[r].cells:
                for p in cell.paragraphs:
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in p.runs:
                        run.font.size = Pt(11)
        doc.add_paragraph()

    # Venue (only if provided)
    if venue.strip():
        venue_para = doc.add_paragraph()
        vrun = venue_para.add_run(f"VENUE : {venue}")
        vrun.font.size = Pt(12)
        venue_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()

    # Duty Table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "SNO."
    hdr[1].text = "Date"
    hdr[2].text = "Team Headed By"
    hdr[3].text = "Team"
    hdr[4].text = "REPORTING TIME"

    for i, duty in enumerate(duties, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = duty["date"].strftime("%d.%m.%Y")
        row[2].text = "\n".join(duty["team_headed"])

        team_lines = []
        for n in duty["p_o"]:
            team_lines.append(f"{n} (P&O)")
        for n in duty["audiologist"]:
            team_lines.append(f"{n} (Audiologist)")
        for n in duty["edp"]:
            team_lines.append(f"{n} (EDP)")
        for n in duty["spectacles"]:
            team_lines.append(f"{n} (Spectacles)")
        for n in duty["technician"]:
            team_lines.append(f"{n} (Technician)")

        row[3].text = "\n".join(team_lines)
        row[4].text = duty["reporting_time"]

    doc.add_paragraph("\n\n")

    # Signatures
    table_sig = doc.add_table(rows=1, cols=2)
    table_sig.autofit = True
    table_sig.style = None
    left_cell = table_sig.rows[0].cells[0]
    right_cell = table_sig.rows[0].cells[1]

    left_para = left_cell.paragraphs[0]
    run1 = left_para.add_run("PRIYANKA SINGH\n")
    run1.font.size = Pt(11)
    run2 = left_para.add_run("AM/P&O")
    run2.bold = True
    run2.font.size = Pt(11)
    left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    right_para = right_cell.paragraphs[0]
    run3 = right_para.add_run("PANKAJ DWIVEDI\n")
    run3.font.size = Pt(11)
    run4 = right_para.add_run("AM&IN-CHARGE")
    run4.bold = True
    run4.font.size = Pt(11)
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ---------------- Export ----------------
st.markdown("### 📥 Export")
if st.button("Generate Duty Chart", type="primary", use_container_width=True):
    if not st.session_state.duties:
        st.warning("⚠️ Please add at least one row first.")
    else:
        docx_io = build_doc(st.session_state.duties, chart_title, venue, sap_id, camp_id, nob, value)
        filename = f"DutyChart_{date.today().strftime('%d%m%Y')}.docx"
        st.success("✅ Duty Chart generated successfully.")
        st.download_button(
            "⬇️ Download Duty Chart (.docx)",
            docx_io,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

